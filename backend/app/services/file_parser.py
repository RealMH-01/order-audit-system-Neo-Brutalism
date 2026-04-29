"""文件解析服务：统一处理文本、PDF、Office 文档和图片输入。"""

from __future__ import annotations

import base64
import io
import logging
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from fastapi import UploadFile

from app.config import Settings
from app.errors import AppError
from app.models.schemas import FeatureStatus, FileCapability, FileRecord
from app.services.runtime_store import RuntimeStore

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover
    load_workbook = None

try:
    from pdf2image import convert_from_bytes
except Exception:  # pragma: no cover
    convert_from_bytes = None

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None

logger = logging.getLogger(__name__)

_DIAG_KEYWORDS = (
    "1.85",
    "1.83",
    "unit price",
    "Unit Price",
    "单价",
    "1,050",
    "1050",
    "20 IBC",
    "38,850",
    "38850",
    "21,000",
    "21000",
    "Contract No",
    "Invoice No",
    "PO No",
)


class FileParserService:
    """管理文件暂存与统一解析，为审核执行层提供标准输入。"""

    def __init__(self, settings: Settings, store: RuntimeStore) -> None:
        self.settings = settings
        self.store = store
        self.max_files_per_user = 10
        self.max_file_size_bytes = 20 * 1024 * 1024

    def get_capability(self) -> FileCapability:
        """返回当前文件解析模块能力说明。"""

        return FileCapability(
            supported_types=["pdf", "docx", "xlsx", "png", "jpg", "jpeg", "txt"],
            features=[
                FeatureStatus(
                    name="统一文件解析入口",
                    ready=True,
                    note="已支持 bytes + filename 解析，并输出统一结构。",
                ),
                FeatureStatus(
                    name="扫描件识别与图像导出",
                    ready=True,
                    note="已提供扫描件 PDF 检测、图片 base64 与 PDF 页图导出结构。",
                ),
            ],
        )

    async def upload_file(self, user_id: str, file: UploadFile) -> FileRecord:
        """上传文件并保存完整解析结果到运行态存储。"""

        self._ensure_user_file_limit(user_id)
        file_bytes = await file.read()
        parsed = self.parse_file(file_bytes, file.filename or "unnamed-file", content_type=file.content_type)
        self.store.files[parsed["id"]] = {**parsed, "user_id": user_id}
        return self._to_file_record(parsed)

    async def replace_file(self, user_id: str, file_id: str, file: UploadFile) -> FileRecord:
        """替换已有文件内容并沿用原 file_id。"""

        self._get_user_file_data(user_id, file_id)
        file_bytes = await file.read()
        parsed = self.parse_file(
            file_bytes,
            file.filename or "unnamed-file",
            content_type=file.content_type,
            file_id=file_id,
        )
        self.store.files[file_id] = {**parsed, "user_id": user_id}
        return self._to_file_record(parsed)

    def delete_file(self, user_id: str, file_id: str) -> None:
        """删除当前用户的暂存文件。"""

        self._get_user_file_data(user_id, file_id)
        del self.store.files[file_id]

    def list_user_files(self, user_id: str) -> list[FileRecord]:
        """列出当前用户仍在运行态暂存的文件。"""

        user_files = [
            item
            for item in self.store.files.values()
            if item.get("user_id") == user_id
        ]
        user_files.sort(key=lambda item: str(item.get("uploaded_at") or ""), reverse=True)
        return [self._to_file_record(item) for item in user_files]

    def delete_files_by_ids(self, user_id: str, file_ids: list[str]) -> int:
        """按 file_id 集合删除当前用户自己的暂存文件。"""

        target_ids = {str(file_id).strip() for file_id in file_ids if str(file_id).strip()}
        deleted_count = 0
        for file_id in list(target_ids):
            file_data = self.store.files.get(file_id)
            if not file_data or file_data.get("user_id") != user_id:
                continue
            del self.store.files[file_id]
            deleted_count += 1
        return deleted_count

    def get_user_file(self, user_id: str, file_id: str) -> dict[str, object]:
        """读取当前用户文件的完整运行态记录。"""

        return self._get_user_file_data(user_id, file_id)

    def parse_file(
        self,
        file_bytes: bytes,
        filename: str,
        *,
        content_type: str | None = None,
        file_id: str | None = None,
    ) -> dict[str, object]:
        """统一解析文件，返回可供 orchestrator 使用的标准结构。"""

        if not file_bytes:
            raise AppError("上传失败，文件内容不能为空。", status_code=400)
        if len(file_bytes) > self.max_file_size_bytes:
            raise AppError("文件过大，当前仅支持 20MB 以内的文件。", status_code=400)

        extension = Path(filename).suffix.lower().lstrip(".")
        detected_type = self._detect_type(filename, extension)
        content_type = content_type or self._guess_content_type(extension)

        parser_map = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "xlsx": self._parse_xlsx,
            "xlsm": self._parse_xlsx,
            "png": self._parse_image,
            "jpg": self._parse_image,
            "jpeg": self._parse_image,
            "txt": self._parse_text,
        }
        parser = parser_map.get(extension, self._parse_text)
        parsed_payload = parser(file_bytes, filename)

        preview_text = self._build_preview_text(parsed_payload.get("text", ""), detected_type)
        keep_raw = (extension == "xlsx") or (bool(parsed_payload.get("needs_ocr")) and not parsed_payload.get("page_images"))
        return {
            "id": file_id or str(uuid4()),
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(file_bytes),
            "detected_type": detected_type,
            "preview_text": preview_text,
            "uploaded_at": datetime.now(timezone.utc),
            "extension": extension,
            "text": parsed_payload.get("text", ""),
            "page_count": parsed_payload.get("page_count", 1),
            "source_kind": parsed_payload.get("source_kind", "text"),
            "parse_mode": parsed_payload.get("parse_mode", "direct"),
            "needs_ocr": bool(parsed_payload.get("needs_ocr", False)),
            "is_scanned_pdf": bool(parsed_payload.get("is_scanned_pdf", False)),
            "page_images": parsed_payload.get("page_images", []),
            "image_base64": parsed_payload.get("image_base64"),
            "raw_bytes": file_bytes if keep_raw else None,
        }

    def _parse_text(self, file_bytes: bytes, _: str) -> dict[str, object]:
        """解析纯文本文件。"""

        text = file_bytes.decode("utf-8", errors="ignore").strip()
        return {
            "text": text,
            "page_count": 1,
            "source_kind": "text",
            "parse_mode": "text",
            "needs_ocr": False,
            "is_scanned_pdf": False,
            "page_images": [],
            "image_base64": None,
        }

    def _parse_docx(self, file_bytes: bytes, filename: str) -> dict[str, object]:
        """解析 Word 文档。"""

        if Document is None:
            raise AppError("当前环境缺少 python-docx，暂时无法解析 Word 文件。", status_code=500)

        document = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        if paragraphs:
            parts.append("--- 文档内容 ---")
            parts.extend(paragraphs)

        if document.tables:
            for t_idx, table in enumerate(document.tables, start=1):
                parts.append(f"\n[表格 {t_idx}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("| " + " | ".join(cells) + " |")

        text = "\n".join(parts) if parts else ""
        self._log_diag_text("docx_parse", text, filename=filename)
        return {
            "text": text,
            "page_count": 1,
            "source_kind": "docx",
            "parse_mode": f"docx:{filename}",
            "needs_ocr": False,
            "is_scanned_pdf": False,
            "page_images": [],
            "image_base64": None,
        }

    def _parse_xlsx(self, file_bytes: bytes, filename: str) -> dict[str, object]:
        """解析 Excel 文档。"""

        if load_workbook is None:
            raise AppError("当前环境缺少 openpyxl，暂时无法解析 Excel 文件。", status_code=500)

        workbook = load_workbook(io.BytesIO(file_bytes), data_only=True)
        chunks: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell not in (None, "")]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                chunks.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
        text = "\n\n".join(chunks).strip()
        self._log_diag_text("xlsx_parse", text, filename=filename)
        return {
            "text": text,
            "page_count": len(workbook.worksheets) or 1,
            "source_kind": "xlsx",
            "parse_mode": "xlsx",
            "needs_ocr": False,
            "is_scanned_pdf": False,
            "page_images": [],
            "image_base64": None,
        }

    def _parse_image(self, file_bytes: bytes, _: str) -> dict[str, object]:
        """解析图片文件，返回 base64 供视觉模型使用。"""

        encoded = base64.b64encode(file_bytes).decode("utf-8")
        return {
            "text": "",
            "page_count": 1,
            "source_kind": "image",
            "parse_mode": "image-base64",
            "needs_ocr": True,
            "is_scanned_pdf": False,
            "page_images": [encoded],
            "image_base64": encoded,
        }

    def _parse_pdf(self, file_bytes: bytes, _: str) -> dict[str, object]:
        """解析 PDF，并在文本不足时标记为扫描件。"""

        extracted_pages: list[str] = []
        page_count = 1
        if pdfplumber is not None:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages) or 1
                for page in pdf.pages:
                    extracted_pages.append((page.extract_text() or "").strip())
        else:
            extracted_pages.append(file_bytes[:400].decode("utf-8", errors="ignore").strip())

        text = "\n\n".join(page for page in extracted_pages if page).strip()
        is_scanned_pdf = len(text) < 80
        page_images = self._pdf_pages_to_base64(file_bytes) if is_scanned_pdf else []

        return {
            "text": text,
            "page_count": page_count,
            "source_kind": "pdf",
            "parse_mode": "pdf-text" if text else "pdf-scan",
            "needs_ocr": is_scanned_pdf,
            "is_scanned_pdf": is_scanned_pdf,
            "page_images": page_images,
            "image_base64": page_images[0] if page_images else None,
        }

    def _pdf_pages_to_base64(self, file_bytes: bytes) -> list[str]:
        """把 PDF 页导出为 PNG base64，供视觉模型或 OCR 判断使用。"""

        if convert_from_bytes is None:
            return []

        images = convert_from_bytes(file_bytes, first_page=1, last_page=3)
        encoded_pages: list[str] = []
        for image in images:
            buffer = io.BytesIO()
            image.save(buffer, format="PNG")
            encoded_pages.append(base64.b64encode(buffer.getvalue()).decode("utf-8"))
        return encoded_pages

    def _ensure_user_file_limit(self, user_id: str) -> None:
        """检查单用户文件数量上限。"""

        file_count = sum(1 for item in self.store.files.values() if item.get("user_id") == user_id)
        if file_count >= self.max_files_per_user:
            raise AppError("单个用户最多暂存 10 个文件，请先删除不需要的文件。", status_code=400)

    def _get_user_file_data(self, user_id: str, file_id: str) -> dict[str, object]:
        """读取并校验当前用户文件访问权限。"""

        file_data = self.store.files.get(file_id)
        if not file_data or file_data.get("user_id") != user_id:
            raise AppError("未找到指定文件，或你无权访问该文件。", status_code=404)
        return file_data

    @staticmethod
    def _detect_type(filename: str, extension: str) -> str:
        """从文件名和扩展名推断单据类型。"""

        name = Path(filename).stem.lower()
        if FileParserService._matches_po_filename(name):
            return "po"
        if FileParserService._matches_invoice_filename(name):
            return "invoice"
        if "packing" in name or "plist" in name:
            return "packing_list"
        if "shipping" in name or "si" in name:
            return "shipping_instruction"
        if "bill_of_lading" in name or "b/l" in name or "bol" in name:
            return "bill_of_lading"
        if "certificate_of_origin" in name or "coo" in name:
            return "certificate_of_origin"
        if "customs" in name or "declaration" in name:
            return "customs_declaration"
        if "letter_of_credit" in name or "_lc" in name or " lc" in name:
            return "letter_of_credit"
        if extension in {"pdf", "docx", "xlsx", "png", "jpg", "jpeg", "txt"}:
            return extension
        return "other"

    @staticmethod
    def _matches_invoice_filename(name: str) -> bool:
        normalized = name.replace(" ", "_").replace(".", "_")
        if "commercial_invoice" in normalized or "invoice" in normalized:
            return True
        if normalized.startswith(("ci-", "ci_", "commercial")):
            return True
        if "-ci-" in normalized or "_ci_" in normalized:
            return True
        return bool(re.search(r"(?:^|[-_\s])inv(?:$|[-_\s])", name, flags=re.IGNORECASE))

    @staticmethod
    def _matches_po_filename(name: str) -> bool:
        normalized = name.replace(" ", "_").replace(".", "_")
        if "purchase_order" in normalized:
            return True
        return bool(re.search(r"(?:^|[-_\s])po(?:$|[-_\s])", name, flags=re.IGNORECASE))

    @staticmethod
    def _collect_diag_hits(text: str, keywords: tuple[str, ...] = _DIAG_KEYWORDS) -> dict[str, int]:
        lowered = text.lower()
        hits: dict[str, int] = {}
        for keyword in keywords:
            count = lowered.count(keyword.lower())
            if count:
                hits[keyword] = count
        return hits

    @staticmethod
    def _collect_diag_snippets(
        text: str,
        keywords: tuple[str, ...] = _DIAG_KEYWORDS,
        *,
        radius: int = 120,
    ) -> dict[str, list[str]]:
        snippets: dict[str, list[str]] = {}
        lowered = text.lower()
        for keyword in keywords:
            index = lowered.find(keyword.lower())
            if index < 0:
                continue
            start = max(0, index - radius)
            end = min(len(text), index + len(keyword) + radius)
            snippets[keyword] = [text[start:end].replace("\n", "\\n")]
        return snippets

    @classmethod
    def _log_diag_text(cls, label: str, text: str, *, filename: str = "") -> None:
        hits = cls._collect_diag_hits(text or "")
        logger.info("DIAG %s filename=%s len=%d keyword_hits=%s", label, filename, len(text or ""), sorted(hits))
        if os.getenv("AUDIT_DEBUG_DIAG", "").lower() == "true":
            logger.info("DIAG %s snippets=%s", label, cls._collect_diag_snippets(text or ""))

    @staticmethod
    def _build_preview_text(text: str, detected_type: str) -> str:
        """构造前端和日志可读的预览文本。"""

        if detected_type in {"png", "jpg", "jpeg"} and not text:
            return "图片文件已接收，当前以视觉/OCR 流程继续处理。"
        preview = (text or "").strip()
        if preview:
            return preview[:300]
        return "文件已接收，但当前没有可直接展示的文本预览。"

    @staticmethod
    def _guess_content_type(extension: str) -> str:
        """根据扩展名补一个基础 content type。"""

        return {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "xlsm": "application/vnd.ms-excel.sheet.macroEnabled.12",
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "txt": "text/plain",
        }.get(extension, "application/octet-stream")

    @staticmethod
    def _to_file_record(parsed: dict[str, object]) -> FileRecord:
        """把完整解析结果压缩成对上传接口友好的响应结构。"""

        return FileRecord(
            id=str(parsed["id"]),
            filename=str(parsed["filename"]),
            content_type=str(parsed["content_type"]),
            size_bytes=int(parsed["size_bytes"]),
            detected_type=str(parsed["detected_type"]),
            preview_text=str(parsed["preview_text"]),
            uploaded_at=parsed["uploaded_at"],
        )
